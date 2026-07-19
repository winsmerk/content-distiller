#!/usr/bin/env python3
"""
Markdown → Word (.docx) 通用转换器
支持：标题、表格（交替行着色）、代码块、引用块、有序/无序列表、
      复选框、粗体/斜体/行内代码、水平线。中文排版优化。

用法：
    from utils.md_to_docx import md_to_docx
    md_to_docx("input.md", "output.docx")
    
    # 命令行
    python md_to_docx.py input.md output.docx
"""

import os
import re
import sys

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ----------------------------------------------------------
# 样式配置（可调整）
# ----------------------------------------------------------
FONT_BODY = "Microsoft YaHei"
FONT_CODE = "Consolas"
FONT_SIZE_BODY = Pt(11)
FONT_SIZE_CODE = Pt(9.5)
FONT_SIZE_TABLE = Pt(9.5)
COLOR_HEADING = RGBColor(0x1A, 0x1A, 0x2E)
COLOR_H2 = RGBColor(0x2C, 0x3E, 0x6B)
COLOR_H3 = RGBColor(0x34, 0x4F, 0x8A)
COLOR_TABLE_HEADER_BG = "2C3E6B"
COLOR_TABLE_ALT_ROW = "F0F4FA"
COLOR_BLOCKQUOTE_BORDER = "2C3E6B"
COLOR_BLOCKQUOTE_BG = "F0F4FA"
COLOR_CODE_BG = "F5F7FA"
COLOR_INLINE_CODE = RGBColor(0xC7, 0x25, 0x4E)


# ----------------------------------------------------------
# 文档样式初始化
# ----------------------------------------------------------
def setup_styles(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_BODY
    font.size = FONT_SIZE_BODY
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.35

    heading_sizes = {1: Pt(22), 2: Pt(18), 3: Pt(14), 4: Pt(12)}
    heading_colors = {1: COLOR_HEADING, 2: COLOR_H2, 3: COLOR_H3, 4: RGBColor(0x45, 0x5A, 0x8E)}
    heading_space_before = {1: Pt(24), 2: Pt(18), 3: Pt(12), 4: Pt(10)}

    for level in range(1, 5):
        sname = f"Heading {level}"
        if sname in doc.styles:
            hs = doc.styles[sname]
            hs.font.name = FONT_BODY
            hs.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
            hs.font.bold = True
            hs.font.size = heading_sizes[level]
            hs.font.color.rgb = heading_colors[level]
            hs.paragraph_format.space_before = heading_space_before[level]
            hs.paragraph_format.space_after = Pt(8 if level <= 2 else 4)


# ----------------------------------------------------------
# 内联格式解析
# ----------------------------------------------------------
def parse_inline(paragraph, text):
    parts = re.split(r"(\*\*.*?\*\*|`[^`]+`|\*[^*]+\*|__.*?__)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("__") and part.endswith("__"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = FONT_CODE
            run.font.size = Pt(10)
            run.font.color.rgb = COLOR_INLINE_CODE
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
            run.element.rPr.append(shading)
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


# ----------------------------------------------------------
# 元素渲染
# ----------------------------------------------------------
def add_table(doc, headers, rows):
    if not headers and not rows:
        return
    num_cols = len(headers) if headers else (len(rows[0]) if rows else 0)
    if num_cols == 0:
        return

    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    if headers:
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header.strip())
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = FONT_BODY
            run.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_shading(cell, COLOR_TABLE_HEADER_BG)

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx >= num_cols:
                break
            cell = table.cell(r_idx + 1, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            parse_inline(p, cell_text.strip())
            for run in p.runs:
                run.font.size = FONT_SIZE_TABLE
                run.font.name = FONT_BODY
                run.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
            if r_idx % 2 == 0:
                set_cell_shading(cell, COLOR_TABLE_ALT_ROW)

    doc.add_paragraph("")


def add_code_block(doc, code_lines):
    for line in code_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(line if line else " ")
        run.font.name = FONT_CODE
        run.font.size = FONT_SIZE_CODE
        run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
        run.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        pPr = p._p.get_or_add_pPr()
        pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{COLOR_CODE_BG}" w:val="clear"/>'))
    doc.add_paragraph("")


def add_blockquote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="12" w:space="8" w:color="{COLOR_BLOCKQUOTE_BORDER}"/>'
        f'</w:pBdr>'
    ))
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{COLOR_BLOCKQUOTE_BG}" w:val="clear"/>'))
    parse_inline(p, text)
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.name = FONT_BODY
        run.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        run.font.color.rgb = RGBColor(0x45, 0x5A, 0x8E)
        run.italic = True


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
        f'</w:pBdr>'
    ))


# ----------------------------------------------------------
# 主转换函数
# ----------------------------------------------------------
def md_to_docx(md_path, docx_path):
    """将 Markdown 文件转换为格式化的 Word 文档"""
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()
    setup_styles(doc)

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    lines = content.split("\n")
    i = 0
    in_code = False
    code_lines = []
    in_table = False
    table_headers = []
    table_rows = []
    table_sep = False

    while i < len(lines):
        line = lines[i]

        # 代码块
        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                if in_table:
                    add_table(doc, table_headers, table_rows)
                    in_table = False
                    table_headers, table_rows, table_sep = [], [], False
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # 表格
        if "|" in line and line.strip().startswith("|") and line.strip().endswith("|"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if all(re.match(r"^[-:]+$", c.strip()) for c in cells if c.strip()):
                table_sep = True
                i += 1
                continue
            if not in_table:
                in_table = True
                table_headers = cells
                table_sep = False
            elif table_sep:
                table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table:
                add_table(doc, table_headers, table_rows)
                in_table = False
                table_headers, table_rows, table_sep = [], [], False

        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # 水平线
        if stripped in ("---", "***", "___"):
            add_horizontal_rule(doc)
            i += 1
            continue

        # 标题
        hm = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if hm:
            level = min(len(hm.group(1)), 4)
            doc.add_heading(hm.group(2).strip(), level=level)
            i += 1
            continue

        # 引用块
        if stripped.startswith(">"):
            qt = stripped[1:].strip()
            while i + 1 < len(lines) and lines[i + 1].strip().startswith(">"):
                i += 1
                qt += "\n" + lines[i].strip()[1:].strip()
            add_blockquote(doc, qt)
            i += 1
            continue

        # 无序列表（注意：对原始 line 匹配以正确检测缩进）
        ul = re.match(r"^(\s*)([-*+])\s+(.+)$", line)
        if ul:
            indent_level = len(ul.group(1)) // 2
            text = ul.group(3)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25 + indent_level * 0.25)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

            if text.startswith("[ ]") or text.startswith("[x]") or text.startswith("[X]"):
                checked = text[1] in ("x", "X")
                p.add_run("\u2611 " if checked else "\u2610 ").font.size = Pt(11)
                text = text[3:].strip()
            else:
                run = p.add_run("\u2022 ")
                run.font.size = Pt(11)
                run.font.color.rgb = COLOR_H2
            parse_inline(p, text)
            i += 1
            continue

        # 有序列表（注意：对原始 line 匹配以正确检测缩进）
        ol = re.match(r"^(\s*)(\d+)[.)]\s+(.+)$", line)
        if ol:
            indent_level = len(ol.group(1)) // 2
            num = ol.group(2)
            text = ol.group(3)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25 + indent_level * 0.25)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f"{num}. ")
            run.bold = True
            run.font.color.rgb = COLOR_H2
            parse_inline(p, text)
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        parse_inline(p, stripped)
        i += 1

    # 收尾
    if in_table:
        add_table(doc, table_headers, table_rows)
    if in_code and code_lines:
        add_code_block(doc, code_lines)

    doc.save(docx_path)
    return docx_path


# ----------------------------------------------------------
# CLI
# ----------------------------------------------------------
if __name__ == "__main__":
    if sys.platform == "win32":
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")

    if len(sys.argv) < 3:
        print("用法: python md_to_docx.py <input.md> <output.docx>")
        sys.exit(1)

    md_path = sys.argv[1]
    docx_path = sys.argv[2]

    if not os.path.exists(md_path):
        print(f"错误: 文件不存在 {md_path}")
        sys.exit(1)

    md_to_docx(md_path, docx_path)
    size_kb = os.path.getsize(docx_path) / 1024
    print(f"✅ {docx_path} ({size_kb:.0f}KB)")
